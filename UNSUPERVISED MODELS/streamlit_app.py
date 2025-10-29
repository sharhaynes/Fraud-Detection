import io
from docx import Document
from docx.shared import Inches
import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.preprocessing import LabelEncoder, StandardScaler
from sklearn.metrics import classification_report, confusion_matrix, roc_auc_score, roc_curve, precision_score, recall_score, f1_score
from sklearn.cluster import KMeans
from sklearn.ensemble import IsolationForest
from sklearn.svm import OneClassSVM
from sklearn.neighbors import LocalOutlierFactor

st.title("Unsupervised Fraud Detection on IPS Data")

train_file = st.file_uploader("Upload Training Dataset (no labels)", type=["csv"])
test_file = st.file_uploader("Upload Testing Dataset (no labels)", type=["csv"])
test_labels_file = st.file_uploader("Upload True Test Labels (isFraud only)", type=["csv"])

def check_and_load(file, name):
    if file is not None:
        try:
            df = pd.read_csv(file)
            st.write(f"{name} dataset preview:")
            st.dataframe(df.head())
            return df
        except Exception as e:
            st.error(f"Error reading {name} file: {e}")
    return None

train_df = check_and_load(train_file, "Training")
test_df = check_and_load(test_file, "Testing")
test_labels_df = check_and_load(test_labels_file, "Test Labels")

le = LabelEncoder()
le_fitted = False

def preprocess(df, fit_scaler=False, scaler=None, fit_label=False):
    global le, le_fitted
    df = df.copy()

    feature_cols = ['type', 'amount', 'oldbalanceOrg', 'newbalanceOrig', 'oldbalanceDest', 'newbalanceDest']

    if fit_label:
        le.fit(df['type'])
        le_fitted = True

    df['type'] = le.transform(df['type']) if le_fitted else df['type']
    X = df[feature_cols]

    if fit_scaler:
        scaler = StandardScaler()
        X_scaled = scaler.fit_transform(X)
    else:
        X_scaled = scaler.transform(X)

    return X_scaled, scaler

if train_df is not None and test_df is not None and test_labels_df is not None:
    X_train, scaler = preprocess(train_df, fit_scaler=True, fit_label=True)
    X_test, _ = preprocess(test_df, fit_scaler=False, scaler=scaler)
    y_test = test_labels_df['isFraud']

    models = {
        "K-Means": KMeans(n_clusters=2, random_state=42),
        "Isolation Forest": IsolationForest(contamination=0.02, random_state=42),
        "One-Class SVM": OneClassSVM(nu=0.02, kernel="rbf", gamma='scale'),
        "Local Outlier Factor": LocalOutlierFactor(n_neighbors=20, contamination=0.02, novelty=True)
    }

    results = {}

    st.subheader("Unsupervised Model Evaluation Results")

    for name, model in models.items():
        if name == "Local Outlier Factor":
            model.fit(X_train)
            preds = model.predict(X_test)
            y_pred = (preds == -1).astype(int)
            scores = -model.decision_function(X_test)
        else:
            model.fit(X_train)
            preds = model.predict(X_test)
            if name == "K-Means":
                counts = np.bincount(preds)
                anomaly_label = np.argmin(counts)
                y_pred = (preds == anomaly_label).astype(int)
                distances = model.transform(X_test)
                scores = distances[:, anomaly_label]
            else:
                y_pred = (preds == -1).astype(int)
                scores = -model.decision_function(X_test)

        precision = precision_score(y_test, y_pred)
        recall = recall_score(y_test, y_pred)
        f1 = f1_score(y_test, y_pred)
        auc = roc_auc_score(y_test, scores)
        cm = confusion_matrix(y_test, y_pred)

        classification_rep = classification_report(y_test, y_pred, output_dict=True)
        macro_avg = classification_rep['macro avg']
        weighted_avg = classification_rep['weighted avg']

        results[name] = {
            "Precision": precision,
            "Recall": recall,
            "F1 Score": f1,
            "AUC": auc,
            "Confusion Matrix": cm,
            "Scores": scores,
            "Predictions": y_pred,
            "Macro Avg": macro_avg,
            "Weighted Avg": weighted_avg
        }

        st.write(f"### {name}")
        st.write("**Classification Report**")
        st.text(classification_report(y_test, y_pred))

        fig_cm, ax_cm = plt.subplots()
        sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', ax=ax_cm)
        ax_cm.set_title(f"{name} - Confusion Matrix")
        ax_cm.set_xlabel("Predicted")
        ax_cm.set_ylabel("Actual")
        st.pyplot(fig_cm)

        fpr, tpr, _ = roc_curve(y_test, scores)
        fig_roc, ax_roc = plt.subplots()
        ax_roc.plot(fpr, tpr, label=f'{name} (AUC = {auc:.2f})')
        ax_roc.plot([0, 1], [0, 1], 'k--')
        ax_roc.set_xlabel('False Positive Rate')
        ax_roc.set_ylabel('True Positive Rate')
        ax_roc.set_title(f"{name} - ROC Curve")
        ax_roc.legend()
        st.pyplot(fig_roc)

    st.subheader("Model Performance Summary")
    summary_data = []
    for model_name, metrics in results.items():
        summary_data.append({
            "Model": model_name,
            "Precision": metrics["Precision"],
            "Recall": metrics["Recall"],
            "F1 Score": metrics["F1 Score"],
            "AUC": metrics["AUC"],
            "Macro Avg F1": metrics["Macro Avg"]["f1-score"],
            "Weighted Avg F1": metrics["Weighted Avg"]["f1-score"]
        })
    summary = pd.DataFrame(summary_data)
    st.dataframe(summary)

    csv_buffer = io.StringIO()
    summary.to_csv(csv_buffer, index=False)
    st.download_button(
        label="📥 Download Summary CSV",
        data=csv_buffer.getvalue(),
        file_name="unsupervised_model_summary.csv",
        mime="text/csv"
    )

else:
    st.info("Please upload all three files: training data, testing data, and test labels.")

def save_word_report(results):
    doc = Document()
    doc.add_heading("IPS Fraud Detection Model Evaluation Report", 0)

    for model_name, metrics in results.items():
        doc.add_heading(model_name, level=1)
        doc.add_paragraph(f"Precision: {metrics['Precision']:.2f}")
        doc.add_paragraph(f"Recall: {metrics['Recall']:.2f}")
        doc.add_paragraph(f"F1 Score: {metrics['F1 Score']:.2f}")
        doc.add_paragraph(f"AUC: {metrics['AUC']:.2f}")

        # Save confusion matrix chart
        fig_cm, ax_cm = plt.subplots()
        sns.heatmap(metrics["Confusion Matrix"], annot=True, fmt='d', cmap='Blues', ax=ax_cm)
        ax_cm.set_title(f"{model_name} - Confusion Matrix")
        cm_path = f"{model_name}_cm.png"
        plt.savefig(cm_path)
        doc.add_picture(cm_path, width=Inches(4))
        plt.close()

        # Save ROC curve chart
        fpr, tpr, _ = roc_curve(y_test, metrics["Scores"])
        fig_roc, ax_roc = plt.subplots()
        ax_roc.plot(fpr, tpr, label=f'{model_name} (AUC = {metrics["AUC"]:.2f})')
        ax_roc.plot([0, 1], [0, 1], 'k--')
        ax_roc.set_xlabel('False Positive Rate')
        ax_roc.set_ylabel('True Positive Rate')
        ax_roc.set_title(f"{model_name} - ROC Curve")
        ax_roc.legend()
        roc_path = f"{model_name}_roc.png"
        plt.savefig(roc_path)
        doc.add_picture(roc_path, width=Inches(4))
        plt.close()

    # Save document
    doc_path = "IPS_Model_Report.docx"
    doc.save(doc_path)
    return doc_path

# Add download link
if st.button("📄 Generate Word Report"):
    report_path = save_word_report(results)
    with open(report_path, "rb") as file:
        st.download_button(
            label="📥 Download Full Word Report",
            data=file,
            file_name="IPS_Model_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )