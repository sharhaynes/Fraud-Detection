import io
import os
from docx import Document
from docx.shared import Inches
import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.preprocessing import LabelEncoder, StandardScaler
from sklearn.metrics import (
    classification_report, confusion_matrix, roc_auc_score, roc_curve,
    precision_score, recall_score, f1_score, average_precision_score, precision_recall_curve
)
from sklearn.semi_supervised import LabelSpreading, LabelPropagation
from sklearn.ensemble import RandomForestClassifier
from sklearn.svm import SVC
from sklearn.semi_supervised import SelfTrainingClassifier

st.title("Semi-Supervised Fraud Detection on IPS Data")

train_file = st.file_uploader("Upload Training Dataset (partially labeled)", type=["csv"])
test_file = st.file_uploader("Upload Testing Dataset", type=["csv"])
test_labels_file = st.file_uploader("Upload True Test Labels (isFraud only)", type=["csv"])

def check_and_load(file, name):
    if file is not None:
        try:
            df = pd.read_csv(file)
            st.write(f"{name} dataset preview:")
            st.dataframe(df.head())
            return df
        except Exception as e:
            st.error(f"Error reading {name} file: {e}")
    return None

train_df = check_and_load(train_file, "Training")
test_df = check_and_load(test_file, "Testing")
test_labels_df = check_and_load(test_labels_file, "Test Labels")

le = LabelEncoder()
def preprocess(df, fit_scaler=False, scaler=None, fit_label=False):
    df = df.copy()
    feature_cols = ['type', 'amount', 'oldbalanceOrg', 'newbalanceOrig', 'oldbalanceDest', 'newbalanceDest']
    if fit_label:
        df['type'] = le.fit_transform(df['type'])
    else:
        df['type'] = le.transform(df['type'])
    X = df[feature_cols]
    if fit_scaler:
        scaler = StandardScaler()
        X_scaled = scaler.fit_transform(X)
    else:
        X_scaled = scaler.transform(X)
    return X_scaled, scaler

if train_df is not None and test_df is not None and test_labels_df is not None:
    y_train = train_df['isFraud'].copy()
    X_train, scaler = preprocess(train_df, fit_scaler=True, fit_label=True)
    X_test, _ = preprocess(test_df, fit_scaler=False, scaler=scaler)
    y_test = test_labels_df['isFraud']

    y_train_semi = np.full_like(y_train, fill_value=-1)
    labeled_indices = y_train.sample(frac=0.05, random_state=42).index
    y_train_semi[labeled_indices] = y_train[labeled_indices]

    models = {
        "Label Spreading": LabelSpreading(kernel='rbf'),
        "Label Propagation": LabelPropagation(kernel='knn'),
        "Self-Training RF": SelfTrainingClassifier(RandomForestClassifier(n_estimators=50)),
        "Self-Training SVM": SelfTrainingClassifier(SVC(probability=True))
    }

    results = {}
    st.subheader("Semi-Supervised Model Evaluation Results")

    for name, model in models.items():
        model.fit(X_train, y_train_semi)
        preds = model.predict(X_test)
        probs = getattr(model, "predict_proba", lambda x: None)(X_test)
        scores = probs[:, 1] if probs is not None else preds

        precision = precision_score(y_test, preds)
        recall = recall_score(y_test, preds)
        f1 = f1_score(y_test, preds)
        auc = roc_auc_score(y_test, scores)
        ap = average_precision_score(y_test, scores)
        cm = confusion_matrix(y_test, preds)

        classification_rep = classification_report(y_test, preds, output_dict=True)
        macro_avg = classification_rep['macro avg']
        weighted_avg = classification_rep['weighted avg']

        results[name] = {
            "Precision": precision,
            "Recall": recall,
            "F1 Score": f1,
            "AUC": auc,
            "AP": ap,
            "Confusion Matrix": cm,
            "Scores": scores,
            "Predictions": preds,
            "Macro Avg": macro_avg,
            "Weighted Avg": weighted_avg
        }

        st.write(f"### {name}")
        st.text(classification_report(y_test, preds))

        fig_cm, ax_cm = plt.subplots()
        sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', ax=ax_cm)
        ax_cm.set_title(f"{name} - Confusion Matrix")
        fig_cm.savefig(f"{name}_cm.png")
        st.pyplot(fig_cm)

        fpr, tpr, _ = roc_curve(y_test, scores)
        fig_roc, ax_roc = plt.subplots()
        ax_roc.plot(fpr, tpr, label=f'{name} (AUC = {auc:.2f})')
        ax_roc.plot([0, 1], [0, 1], 'k--')
        ax_roc.set_title(f"{name} - ROC Curve")
        ax_roc.set_xlabel("False Positive Rate")
        ax_roc.set_ylabel("True Positive Rate")
        ax_roc.legend()
        fig_roc.savefig(f"{name}_roc.png")
        st.pyplot(fig_roc)

        prec, rec, _ = precision_recall_curve(y_test, scores)
        fig_pr, ax_pr = plt.subplots()
        ax_pr.plot(rec, prec, label=f'{name} (AP = {ap:.2f})')
        ax_pr.set_title(f"{name} - Precision-Recall Curve")
        ax_pr.set_xlabel("Recall")
        ax_pr.set_ylabel("Precision")
        ax_pr.legend()
        fig_pr.savefig(f"{name}_pr.png")
        st.pyplot(fig_pr)

    st.subheader("Model Performance Summary")
    summary_data = []
    for model_name, metrics in results.items():
        summary_data.append({
            "Model": model_name,
            "Precision": metrics["Precision"],
            "Recall": metrics["Recall"],
            "F1 Score": metrics["F1 Score"],
            "AUC": metrics["AUC"],
            "AP": metrics["AP"],
            "Macro Avg F1": metrics["Macro Avg"]["f1-score"],
            "Weighted Avg F1": metrics["Weighted Avg"]["f1-score"]
        })
    summary = pd.DataFrame(summary_data)
    st.dataframe(summary)

    csv_buffer = io.StringIO()
    summary.to_csv(csv_buffer, index=False)
    st.download_button(
        label="📥 Download Summary CSV",
        data=csv_buffer.getvalue(),
        file_name="semi_supervised_model_summary.csv",
        mime="text/csv"
    )

    def save_word_report(results):
        doc = Document()
        doc.add_heading("IPS Fraud Detection - Semi-Supervised Model Report", 0)

        for model_name, metrics in results.items():
            doc.add_heading(model_name, level=1)
            doc.add_paragraph(f"Precision: {metrics['Precision']:.2f}")
            doc.add_paragraph(f"Recall: {metrics['Recall']:.2f}")
            doc.add_paragraph(f"F1 Score: {metrics['F1 Score']:.2f}")
            doc.add_paragraph(f"AUC: {metrics['AUC']:.2f}")
            doc.add_paragraph(f"Average Precision (PR AUC): {metrics['AP']:.2f}")

            for suffix in ["cm", "roc", "pr"]:
                image_path = f"{model_name}_{suffix}.png"
                if os.path.exists(image_path):
                    doc.add_picture(image_path, width=Inches(5))

        doc_path = "IPS_SemiSupervised_Report.docx"
        doc.save(doc_path)
        return doc_path

    if st.button("📄 Generate Word Report"):
        report_path = save_word_report(results)
        with open(report_path, "rb") as f:
            st.download_button(
                label="📥 Download Word Report",
                data=f,
                file_name="IPS_SemiSupervised_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
